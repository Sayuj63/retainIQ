"""
RetainIQ Pro — Production-Level PRD Generator
Generates RetainIQ-Pro-PRD.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Brand colors ──────────────────────────────────────────────────────────────
BRAND_PURPLE = RGBColor(0x3C, 0x34, 0x89)
BRAND_LIGHT  = RGBColor(0xEE, 0xED, 0xFE)
DARK_TEXT    = RGBColor(0x11, 0x11, 0x1A)
MID_TEXT     = RGBColor(0x44, 0x44, 0x55)
LIGHT_BG     = RGBColor(0xF7, 0xF7, 0xFF)
RED_ACCENT   = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x27, 0x80, 0x4F)
AMBER_ACCENT = RGBColor(0x97, 0x67, 0x14)
TABLE_HEADER = RGBColor(0x3C, 0x34, 0x89)
TABLE_ALT    = RGBColor(0xF4, 0xF3, 0xFF)
BORDER_COLOR = RGBColor(0xCC, 0xC9, 0xF0)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **borders):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in borders.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color: RGBColor = None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    add_run(p, text, bold=True, size=20, color=BRAND_PURPLE, font="Calibri")
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), "3C3489")
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, text, bold=True, size=14, color=DARK_TEXT, font="Calibri")
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    add_run(p, text, bold=True, size=12, color=BRAND_PURPLE, font="Calibri")
    return p

def h4(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    add_run(p, text, bold=True, size=11, color=MID_TEXT, font="Calibri")
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, text, size=11, color=DARK_TEXT)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(1)
    if bold_prefix:
        add_run(p, bold_prefix + ": ", bold=True, size=11, color=DARK_TEXT)
        add_run(p, text, size=11, color=DARK_TEXT)
    else:
        add_run(p, text, size=11, color=DARK_TEXT)
    return p

def callout(title, text, accent: RGBColor = None, bg_hex="F4F3FF", border_hex="3C3489"):
    """Single-cell table used as a callout box."""
    color = accent or BRAND_PURPLE
    tbl   = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    set_cell_border(cell, left=border_hex)
    p = cell.paragraphs[0]
    if title:
        add_run(p, title + "  ", bold=True, size=11, color=color)
    add_run(p, text, size=11, color=DARK_TEXT)
    doc.add_paragraph()

def table_2col(rows_data, col_widths=(Cm(5), Cm(11)), header=True):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (a, b) in enumerate(rows_data):
        r = tbl.rows[i]
        r.cells[0].width = col_widths[0]
        r.cells[1].width = col_widths[1]
        if i == 0 and header:
            set_cell_bg(r.cells[0], "3C3489")
            set_cell_bg(r.cells[1], "3C3489")
            p0 = r.cells[0].paragraphs[0]
            p1 = r.cells[1].paragraphs[0]
            add_run(p0, a, bold=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF))
            add_run(p1, b, bold=True, size=11, color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            bg = "F4F3FF" if i % 2 == 0 else "FFFFFF"
            set_cell_bg(r.cells[0], bg)
            set_cell_bg(r.cells[1], bg)
            p0 = r.cells[0].paragraphs[0]
            p1 = r.cells[1].paragraphs[0]
            add_run(p0, a, bold=True, size=10, color=MID_TEXT)
            add_run(p1, b, size=10, color=DARK_TEXT)
    doc.add_paragraph()
    return tbl

def table_generic(headers, rows, col_widths=None):
    all_rows = [headers] + rows
    tbl = doc.add_table(rows=len(all_rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row_data in enumerate(all_rows):
        r = tbl.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = r.cells[j]
            if col_widths:
                cell.width = col_widths[j]
            if i == 0:
                set_cell_bg(cell, "3C3489")
                p = cell.paragraphs[0]
                add_run(p, str(cell_text), bold=True, size=10, color=RGBColor(0xFF,0xFF,0xFF))
            else:
                bg = "F4F3FF" if i % 2 == 1 else "FFFFFF"
                set_cell_bg(cell, bg)
                p = cell.paragraphs[0]
                add_run(p, str(cell_text), size=10, color=DARK_TEXT)
    doc.add_paragraph()
    return tbl

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run()
    run.add_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
add_run(cover, "RetainIQ Pro", bold=True, size=42, color=BRAND_PURPLE, font="Calibri")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, "Product Requirements Document", size=20, color=MID_TEXT)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub2, "Post-Purchase Retention Intelligence Platform for Shopify D2C Brands",
        italic=True, size=13, color=MID_TEXT)

doc.add_paragraph()
doc.add_paragraph()

meta_data = [
    ("Version",     "1.0 — Production Release"),
    ("Date",        datetime.date.today().strftime("%B %d, %Y")),
    ("Status",      "APPROVED — Ready for Engineering"),
    ("Author",      "Product Team"),
    ("Confidentiality", "Internal — Do Not Distribute"),
]
meta_tbl = doc.add_table(rows=len(meta_data), cols=2)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta_data):
    meta_tbl.rows[i].cells[0].width = Cm(5)
    meta_tbl.rows[i].cells[1].width = Cm(9)
    set_cell_bg(meta_tbl.rows[i].cells[0], "EEEDFE")
    set_cell_bg(meta_tbl.rows[i].cells[1], "FFFFFF")
    add_run(meta_tbl.rows[i].cells[0].paragraphs[0], k, bold=True, size=10, color=BRAND_PURPLE)
    add_run(meta_tbl.rows[i].cells[1].paragraphs[0], v, size=10, color=DARK_TEXT)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary")

body(
    "RetainIQ Pro is a production-grade, AI-powered post-purchase retention platform purpose-built for "
    "Shopify Direct-to-Consumer (D2C) brands. It closes the single most expensive gap in modern e-commerce: "
    "brands spend 5–10× more acquiring a customer than keeping one, yet 80 % of post-purchase journeys are "
    "handled by generic confirmation emails that drive zero incremental LTV."
)
body(
    "RetainIQ Pro turns every completed order into a structured, data-rich retention event. The platform "
    "automatically scores churn risk within two hours of purchase, fires personalised WhatsApp / SMS / email "
    "sequences timed to each customer's behaviour, predicts individual SKU replenishment dates with "
    "ML-grade accuracy, and closes the loop with review and UGC capture at the exact moment satisfaction "
    "peaks — not 30 days later when sentiment has faded."
)

h2("Key Outcomes (validated beta data)")
table_generic(
    ["Metric", "Lift vs. baseline", "Measurement window"],
    [
        ["Repeat purchase rate",      "+34 %",      "6 months post-install"],
        ["Customer LTV",              "2.1×",       "6 months post-install"],
        ["Average review volume",     "+290 %",     "Per order, 72 h window"],
        ["Return rate",               "−18 %",      "30 days post-delivery"],
        ["WhatsApp open rate",        "87 %",       "vs. 24 % email baseline"],
        ["Time to first repurchase",  "−11 days",   "Median, consumables category"],
    ],
    col_widths=[Cm(6), Cm(4), Cm(6)]
)

body("Average Shopify merchant setup time: < 1 hour (plug-and-play OAuth flow).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT & MARKET OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Problem Statement & Market Opportunity")

h2("2.1 The Retention Crisis in D2C")
body(
    "The D2C e-commerce sector has undergone a structural shift since iOS 14.5 privacy changes in 2021 "
    "decimated Meta ROAS benchmarks. Customer acquisition costs rose 60 % between 2021–2024, while "
    "average repeat-purchase rates for Shopify brands under $10M ARR remain below 28 %. Brands are "
    "running cash-flow-negative acquisition engines subsidised by first orders that never convert to "
    "long-term customers."
)

h2("2.2 Root Causes")
bullet("No structured post-purchase journey — the average D2C brand sends 1.2 emails after an order, all logistics-focused.")
bullet("Zero behavioural intelligence — repurchase nudges go out on fixed 30-day schedules regardless of product type, AOV, or customer behaviour.")
bullet("Channel mismatch — email open rates for promotional content average 18–24 %; the target audience prefers WhatsApp (87 % open rate) but most brands lack WhatsApp Business API access.")
bullet("Review timing failure — review requests sent after 30 days see <4 % conversion; sent at 72 hours post-delivery conversion exceeds 22 %.")
bullet("No churn signal — brands discover a customer has churned 90+ days after the fact, when win-back is nearly impossible.")

h2("2.3 Market Opportunity")
table_generic(
    ["Segment", "Size (2025)", "CAGR", "Addressable?"],
    [
        ["Global D2C e-commerce market",          "$239 B",   "14.3 %", "Partially"],
        ["Shopify merchant base (active stores)",  "4.6 M",    "—",      "Yes — via app store"],
        ["Shopify Plus merchants (enterprise)",    "~32 000",  "—",      "Yes — direct sales"],
        ["Customer retention software market",     "$8.1 B",   "16.1 %", "Yes — horizontal"],
        ["WhatsApp Commerce TAM",                  "$3.6 B",   "27 %",   "Yes — channel layer"],
    ],
    col_widths=[Cm(6.5), Cm(3), Cm(2.5), Cm(3)]
)

h2("2.4 Competitive Gap")
body(
    "Existing solutions (Klaviyo, Yotpo, Gorgias, AfterShip) each own one layer of the stack. "
    "No single platform unifies behavioural scoring + omnichannel sequencing + replenishment AI + "
    "review capture in a single Shopify-native app. RetainIQ Pro is that platform."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. TARGET MARKET & USER PERSONAS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Target Market & User Personas")

h2("3.1 Primary Segment")
body("Shopify D2C brands generating $500K – $15M ARR in consumables, beauty, wellness, pet, and food categories with repeat-purchase potential.")

h2("3.2 Personas")

h3("Persona A — The Growth-Obsessed Founder (\"Maya\")")
table_2col([
    ["Attribute", "Detail"],
    ["Role",          "Founder / CEO, 8-person brand"],
    ["Revenue",       "$1.2M ARR, 60 % from new customer acquisition"],
    ["Pain",          "Meta ROAS down to 1.4×; can't afford to keep buying new customers"],
    ["Goal",          "Double repeat rate from 22 % → 44 % in 12 months"],
    ["Tech stack",    "Shopify, Klaviyo, WhatsApp Business (manual), no dedicated CRM"],
    ["Key need",      "Setup in under 1 hour; no engineering resources; clear ROI dashboard"],
    ["Willingness",   "Pays $299/mo if it demonstrably reduces CAC payback period"],
])

h3("Persona B — The CRM Manager (\"Rohan\")")
table_2col([
    ["Attribute", "Detail"],
    ["Role",          "Head of CRM / Retention, 45-person brand"],
    ["Revenue",       "$8M ARR, already has Klaviyo + Yotpo"],
    ["Pain",          "Klaviyo automations are channel-siloed; WhatsApp is a manual side project"],
    ["Goal",          "Unified omnichannel flows, predictive replenishment, reduced tool sprawl"],
    ["Tech stack",    "Shopify Plus, Klaviyo, Yotpo, Postscript, AfterShip — 5 tools doing 1 job"],
    ["Key need",      "Native WhatsApp + email + SMS orchestration from one place; segment export"],
    ["Willingness",   "Pays $799/mo for platform consolidation + advanced analytics"],
])

h3("Persona C — The Shopify Plus Enterprise Brand (\"NykaaFashion-tier\")")
table_2col([
    ["Attribute", "Detail"],
    ["Role",          "VP E-commerce + dedicated data team"],
    ["Revenue",       "$15M+ ARR"],
    ["Pain",          "No predictive layer; A/B testing is manual; no ML for replenishment"],
    ["Goal",          "Custom ML models, data warehouse export, SLA-backed uptime, dedicated CSM"],
    ["Tech stack",    "Shopify Plus, custom ERP, BigQuery"],
    ["Key need",      "API-first integration, SSO, raw event streaming, custom model endpoints"],
    ["Willingness",   "Pays $2500+/mo on annual contract"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PRODUCT VISION & STRATEGIC GOALS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Product Vision & Strategic Goals")

h2("4.1 Vision Statement")
callout(
    "",
    "\"Make every post-purchase moment a compounding asset — not a missed opportunity. "
    "RetainIQ Pro is the operating system for D2C retention: behavioural intelligence, "
    "personalised omnichannel sequencing, and predictive reorder — all from one Shopify app.\"",
)

h2("4.2 Strategic Goals (12-month)")
table_generic(
    ["Goal", "Target", "Success Metric"],
    [
        ["G1 — Adoption",          "2 000 active Shopify installs",                "Monthly active stores"],
        ["G2 — Revenue",           "$2.4M ARR by month 12",                        "MRR × 12"],
        ["G3 — Retention",         ">92 % net revenue retention",                  "NRR rolling 6-month"],
        ["G4 — Merchant outcome",  "+30 % avg repeat rate lift for cohort",        "Repeat rate delta 6-month"],
        ["G5 — Channel coverage",  "WhatsApp live in India, Indonesia, Brazil",    "Geo-active WA merchants"],
        ["G6 — Platform trust",    "SOC 2 Type II + GDPR + DPDP (India) compliant","Audit completion date"],
    ],
    col_widths=[Cm(4), Cm(5.5), Cm(5.5)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CORE FEATURES — FULL SPECIFICATION
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Core Features — Full Specification")

# ── 5.1 Behavioral Scoring ────────────────────────────────────────────────────
h2("5.1 Behavioral Churn-Risk Scoring Engine")
body(
    "Assigns a real-time churn-risk score (0–100) to every customer within 2 hours of order "
    "completion. Score drives downstream automation: high-risk customers enter a win-back sequence; "
    "low-risk customers enter a referral / upsell sequence."
)

h3("5.1.1 Input Signals")
table_generic(
    ["Signal", "Weight", "Source", "Availability"],
    [
        ["AOV vs. brand median",              "High",   "Shopify Order API",      "Real-time"],
        ["Number of previously viewed SKUs",  "High",   "Pixel / Storefront API", "Real-time"],
        ["Cart abandonment history",          "High",   "Shopify Cart API",       "Historical"],
        ["Days since last purchase",          "Medium", "Customer API",           "Historical"],
        ["Discount code used (Y/N)",          "Medium", "Order API",              "Real-time"],
        ["Product category churn baseline",   "Medium", "Internal benchmark DB",  "Pre-computed"],
        ["NPS score (if prior customer)",     "Medium", "Review module",          "Historical"],
        ["Geographic region",                 "Low",    "Order shipping address",  "Real-time"],
        ["Device type at order",              "Low",    "Pixel",                  "Real-time"],
        ["Email open history (last 90 days)", "Low",    "Email provider webhook",  "Historical"],
    ],
    col_widths=[Cm(5), Cm(2), Cm(4.5), Cm(2.5)]
)

h3("5.1.2 Model Architecture")
bullet("Algorithm: Gradient-boosted decision trees (LightGBM) + online learning layer for per-brand fine-tuning")
bullet("Training cadence: Weekly global retrain; daily incremental update per brand")
bullet("Cold-start: Rule-based scoring (AOV + category baseline) for brands with < 200 orders")
bullet("Score freshness SLA: Score available within 90 seconds of order.paid webhook receipt")
bullet("Score API: REST endpoint GET /v1/customers/{id}/churn-score returns score + top-3 contributing features")

h3("5.1.3 Acceptance Criteria")
bullet("AC-1: Score produced for 100 % of orders within 120 s of order.paid event under p99 load")
bullet("AC-2: Precision @ threshold 70 ≥ 0.72 on held-out validation set")
bullet("AC-3: Score API latency p50 < 80 ms, p99 < 300 ms")
bullet("AC-4: Model drift alert fires if AUC drops > 0.05 from baseline on rolling 7-day window")
bullet("AC-5: Merchant can override score manually via dashboard for any customer")

# ── 5.2 Smart Sequencing ─────────────────────────────────────────────────────
h2("5.2 Smart Sequencing & Omnichannel Automation")
body(
    "An event-driven orchestration engine that auto-triggers personalised WhatsApp, SMS, and email "
    "touchpoints along the post-purchase journey. All channel selection, timing, and content are "
    "dynamically personalised per customer."
)

h3("5.2.1 Journey Triggers & Timing")
table_generic(
    ["Trigger", "Channel priority", "Content type", "Condition"],
    [
        ["order.paid (T+0)",       "WhatsApp > Email", "Branded order confirmation + unboxing preview",  "Always"],
        ["T+2h (churn score)",     "Internal only",    "Score computed → branch logic activated",         "Always"],
        ["T+2h (high risk ≥70)",   "WhatsApp > SMS",   "Personalised discount offer (dynamic %)",        "Score ≥ 70"],
        ["T+2h (low risk <40)",    "WhatsApp > Email", "Referral programme invite",                       "Score < 40"],
        ["T+24h",                  "WhatsApp > SMS",   "Usage tip + community invite",                    "Always"],
        ["T+72h",                  "WhatsApp > Email", "Review + UGC capture request",                   "Delivered=true"],
        ["T+Day N (replenishment)","WhatsApp > SMS",   "1-tap reorder nudge",                             "Consumable SKU"],
        ["T+Day N+7 (no reorder)", "SMS > Email",      "Win-back with FOMO offer",                        "No reorder placed"],
        ["T+60 days (dormant)",    "Email",            "Re-engagement flow with best-seller showcase",    "0 orders in 60d"],
    ],
    col_widths=[Cm(3.5), Cm(3), Cm(5), Cm(3)]
)

h3("5.2.2 Channel Engine Specs")
bullet("WhatsApp", bold_prefix="Provider")
bullet("Meta WhatsApp Business Cloud API (primary); Gupshup fallback for India tier", level=1)
bullet("Template message approval management built into admin dashboard", level=1)
bullet("Delivery receipt webhook → retry logic (max 3 attempts, 30-min intervals)", level=1)
bullet("SMS", bold_prefix="Provider")
bullet("Twilio (global); MSG91 (India); Vonage (EU)", level=1)
bullet("Carrier lookup to avoid landline delivery waste", level=1)
bullet("Email", bold_prefix="Provider")
bullet("SendGrid (transactional); Klaviyo list-sync for campaign overlap prevention", level=1)
bullet("DKIM + SPF + DMARC enforced on all outbound domains", level=1)

h3("5.2.3 Personalisation Variables")
table_2col([
    ["Variable", "Source"],
    ["{{first_name}}",           "Shopify Customer API"],
    ["{{product_name}}",         "Order line item"],
    ["{{discount_pct}}",         "AI discount engine (risk-tiered)"],
    ["{{reorder_date}}",         "Replenishment model output"],
    ["{{referral_link}}",        "Generated per-customer referral URL"],
    ["{{review_link}}",          "Dynamic deep-link to review widget"],
    ["{{usage_tip}}",            "SKU-level tip library (brand-uploaded)"],
    ["{{community_link}}",       "Brand WhatsApp group / Discord"],
])

h3("5.2.4 Acceptance Criteria")
bullet("AC-1: Message delivery within 5 min of trigger time for ≥ 99 % of events under normal load")
bullet("AC-2: Channel fallback executes within 10 min if primary delivery fails")
bullet("AC-3: Unsubscribe / opt-out processed within 30 s; no further messages sent")
bullet("AC-4: Personalisation variables render correctly for 100 % of messages (no {{unfilled}} tokens)")
bullet("AC-5: Merchant A/B test (content/timing) reduces to < 20 % traffic split with stat-sig detection at p < 0.05")

# ── 5.3 Replenishment AI ─────────────────────────────────────────────────────
h2("5.3 Replenishment AI")
body(
    "Predicts the exact date each customer will run out of a given SKU, triggers a personalised "
    "reorder nudge 3 days before predicted depletion, and enables one-tap reorder via WhatsApp."
)

h3("5.3.1 Prediction Model")
table_2col([
    ["Attribute", "Specification"],
    ["Algorithm",         "Bayesian survival model (Weibull-gamma mixture) per customer×SKU pair"],
    ["Features",          "Purchase cadence history, SKU volume/weight, product type taxonomy, household size proxy (AOV), subscription flag"],
    ["Cold start",        "Category-level median from brand's order history; global prior for new brands"],
    ["Minimum data",      "2 prior purchases of same SKU to activate personalised prediction"],
    ["Confidence output", "Predicted date ± uncertainty interval (shown to merchant, not customer)"],
    ["Model update",      "Updated on every new purchase event in real time"],
    ["Accuracy KPI",      "Predicted vs actual reorder date MAE ≤ 5 days on held-out cohort"],
])

h3("5.3.2 WhatsApp One-Tap Reorder Flow")
bullet("Customer receives WhatsApp message 3 days before predicted depletion date")
bullet("Message contains product image, name, price, and a single CTA: \"Reorder Now\"")
bullet("Tap creates a pre-filled Shopify checkout URL with saved payment method (Shop Pay)")
bullet("Order confirmation echoed back to WhatsApp thread within 60 s of purchase")
bullet("If no reorder in 7 days: escalated win-back sequence triggers")

h3("5.3.3 Merchant Controls")
bullet("Set global or per-SKU \"advance notice\" days (default: 3, range: 1–14)")
bullet("Exclude SKUs from replenishment programme")
bullet("Override predicted date manually per customer")
bullet("View replenishment forecast calendar (next 30/60/90 days)")

h3("5.3.4 Acceptance Criteria")
bullet("AC-1: Replenishment date computed for all eligible customer×SKU pairs within 5 min of qualifying 2nd purchase")
bullet("AC-2: Reorder conversion rate from replenishment WhatsApp ≥ 18 % (validated in beta)")
bullet("AC-3: One-tap reorder checkout load time < 1.5 s on 4G connection")
bullet("AC-4: MAE ≤ 5 days for consumables category at brand with ≥ 500 orders")

# ── 5.4 Review Loop ───────────────────────────────────────────────────────────
h2("5.4 Review & UGC Capture Loop")
body(
    "Captures star ratings, written reviews, photo/video UGC, and NPS scores at T+72h post-delivery "
    "(peak satisfaction). Automatically routes responses: 5-star → Google/Meta amplification; "
    "3-star → support ticket; 1-2 star → private response + recovery offer."
)

h3("5.4.1 Capture Flow")
bullet("T+72h WhatsApp message: \"How's [product]? Rate in one tap\" (emoji buttons: ⭐–⭐⭐⭐⭐⭐)")
bullet("5-star tap: prompts for written review + optional photo/video upload (inline WA flow)")
bullet("3-star tap: soft-surveys for pain point (sizing / quality / delivery) → routes to support")
bullet("1-2 star tap: immediate human-agent alert + recovery offer (replacement or refund)")
bullet("Email fallback for non-WhatsApp customers: embedded star-rating widget (no redirect required)")

h3("5.4.2 Review Routing & Amplification")
table_generic(
    ["Rating", "Action", "Platform", "Delay"],
    [
        ["5★",   "Auto-publish + share link",          "Google Business + Meta (Facebook/Instagram)", "Immediate"],
        ["4★",   "Publish + soft upsell prompt",       "Shopify storefront reviews widget",           "Immediate"],
        ["3★",   "Hold + support ticket",               "Internal helpdesk (Gorgias / Freshdesk)",    "< 5 min"],
        ["1-2★", "Private + recovery offer sent",      "WhatsApp direct message",                     "< 2 min"],
    ],
    col_widths=[Cm(1.5), Cm(5.5), Cm(5), Cm(2)]
)

h3("5.4.3 UGC Management")
bullet("Photo/video stored in brand's RetainIQ media library (S3-backed, CDN-served)")
bullet("Merchant can approve / reject UGC before it appears on storefront")
bullet("One-click republish to Instagram Stories via Meta Content Publishing API")
bullet("UGC tagged with product SKU, review date, and customer segment for filtering")

h3("5.4.4 NPS Engine")
bullet("Separate NPS survey fires at T+14d post-delivery (decoupled from review flow)")
bullet("Detractors (0–6): auto-tagged in CRM with 'at-risk'; personal follow-up queued")
bullet("Promoters (9–10): invited to referral programme automatically")
bullet("NPS trended over time in analytics dashboard with segment drill-down")

h3("5.4.5 Acceptance Criteria")
bullet("AC-1: Review request delivered within 5 min of T+72h trigger for ≥ 99.5 % of eligible orders")
bullet("AC-2: 5-star reviews published to Google/Meta within 2 min of submission")
bullet("AC-3: 1-2 star alert delivered to merchant within 2 min; no auto-publish")
bullet("AC-4: UGC upload supports JPEG/PNG/MP4; max 50 MB; transcoded to CDN within 3 min")
bullet("AC-5: Review collection rate (responses / requests) ≥ 22 % as baseline (validated beta)")

# ── 5.5 Branded Tracking Page ────────────────────────────────────────────────
h2("5.5 Branded Tracking Experience")
body(
    "Replaces Shopify's generic order confirmation and carrier tracking pages with a fully "
    "branded, conversion-optimised experience that drives the next purchase while the customer "
    "is at peak excitement."
)

h3("5.5.1 Page Components")
table_2col([
    ["Component", "Description"],
    ["Order summary",        "Real-time carrier tracking widget (EasyPost API) with visual progress bar"],
    ["Unboxing preview",     "Brand-uploaded GIF / short video of product unboxing experience"],
    ["'What to expect' rail","Customisable content cards: usage tips, setup guides, community links"],
    ["Next-order module",    "Personalised bestseller recommendation (based on purchase + browse history)"],
    ["Referral CTA",         "Unique referral link with social share to WhatsApp / Instagram"],
    ["Review CTA (T+72h)",   "Star rating widget appears inline once 72h post-delivery passes"],
    ["Loyalty progress bar", "Points balance and next-tier progress (if loyalty enabled)"],
])

h3("5.5.2 Technical Specs")
bullet("Hosted at trackingiq.app/{brand_handle}/orders/{order_token} — fully branded subdomain support")
bullet("Page load: LCP < 2.0 s on 4G; CLS < 0.1; FID < 100 ms (Core Web Vitals compliant)")
bullet("Carrier support: DHL, FedEx, UPS, USPS, Delhivery, Bluedart, Shiprocket, 500+ via EasyPost")
bullet("Mobile-first responsive; PWA-installable via \"Add to Home Screen\" prompt")
bullet("Analytics: every element click tracked → feeds back into next-session personalisation")

# ── 5.6 Analytics Dashboard ──────────────────────────────────────────────────
h2("5.6 Analytics & Intelligence Dashboard")

h3("5.6.1 Merchant Dashboard Modules")
table_2col([
    ["Module", "Contents"],
    ["Retention Overview",   "Repeat rate trend, LTV curve, churn rate, cohort heatmap (monthly)"],
    ["Flow Performance",     "Per-sequence open / click / conversion rates; revenue attributed per flow"],
    ["Channel Analytics",    "WhatsApp / SMS / Email side-by-side delivery, open, click, conversion"],
    ["Replenishment Forecast","30/60/90 day reorder volume forecast; individual customer reorder calendar"],
    ["Review Intelligence",  "Star distribution, NPS trend, UGC library, review velocity by SKU"],
    ["Customer Segments",    "AI-generated segments (Champions / Loyal / At-Risk / Dormant / New)"],
    ["Revenue Attribution",  "RetainIQ-attributed revenue per channel per time period"],
    ["A/B Test Results",     "Live and completed experiments with confidence intervals and winner calls"],
])

h3("5.6.2 Data Export & Integrations")
bullet("CSV / JSON export of all metrics on any time range")
bullet("BigQuery / Snowflake live sync (Growth and Enterprise tiers)")
bullet("Klaviyo segment sync: push RetainIQ segments as Klaviyo lists in real time")
bullet("Webhook API: POST to any URL on score-change, review-received, reorder-placed events")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. TECHNICAL ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Technical Architecture")

h2("6.1 System Overview")
body(
    "RetainIQ Pro is built as a cloud-native, multi-tenant SaaS application on AWS. "
    "The architecture is microservices-based with event-driven communication via Apache Kafka. "
    "All services are containerised (Docker + EKS) and deployed via GitOps (ArgoCD)."
)

h2("6.2 Core Services")
table_generic(
    ["Service", "Responsibility", "Tech Stack", "Scaling"],
    [
        ["ingestion-svc",      "Shopify webhook receiver, event validation, Kafka producer",                  "Node.js / Fastify",         "Horizontal, 20 pods max"],
        ["scoring-svc",        "Churn-risk score computation, model serving",                                 "Python / FastAPI + LightGBM","GPU-backed, auto-scale"],
        ["orchestration-svc",  "Journey trigger logic, branch evaluation, message scheduling",                "Node.js / BullMQ",          "Horizontal, queue-backed"],
        ["channel-svc",        "WhatsApp / SMS / Email dispatch, delivery tracking, retry logic",             "Node.js / Fastify",         "Horizontal, per-channel pod"],
        ["replenishment-svc",  "Bayesian survival model inference, reorder date computation",                 "Python / FastAPI",          "CPU, auto-scale"],
        ["review-svc",         "Review capture, routing logic, UGC storage, Google/Meta publish",             "Node.js / Fastify",         "Horizontal"],
        ["tracking-svc",       "Branded tracking page serving, EasyPost polling, cache invalidation",         "Next.js / Vercel Edge",     "Edge CDN"],
        ["analytics-svc",      "Metrics aggregation, dashboard data, cohort computation",                     "Python / FastAPI + ClickHouse","Read replicas"],
        ["api-gateway",        "Auth (JWT + OAuth), rate limiting, routing, CORS",                            "Kong Gateway",              "HA, 3 replicas min"],
        ["admin-svc",          "Merchant onboarding, settings, billing, Shopify OAuth",                       "Next.js 14 + tRPC",         "Serverless"],
    ],
    col_widths=[Cm(3.5), Cm(5.5), Cm(3), Cm(2)]
)

h2("6.3 Data Infrastructure")
table_2col([
    ["Layer", "Technology"],
    ["Primary DB",         "PostgreSQL 16 (RDS Multi-AZ) — transactional data, customer records, flow configs"],
    ["Event store",        "Apache Kafka (MSK) — all webhook events, state transitions, audit trail"],
    ["ML feature store",   "Redis 7 (ElastiCache) — real-time feature vectors, score cache"],
    ["Analytics store",    "ClickHouse (self-hosted EKS) — time-series metrics, cohort aggregations"],
    ["Object storage",     "S3 + CloudFront CDN — UGC assets, tracking page static assets"],
    ["Search",             "OpenSearch — customer search, SKU lookup in merchant dashboard"],
    ["Cache",              "Redis (session, API response cache, idempotency keys)"],
    ["Secrets",            "AWS Secrets Manager — API keys, WhatsApp tokens, SendGrid credentials"],
])

h2("6.4 Shopify Integration Architecture")
bullet("OAuth 2.0 install flow — merchant clicks 'Install' in Shopify App Store → OAuth handshake → access token stored encrypted in Secrets Manager")
bullet("Webhooks registered on install: orders/paid, orders/fulfilled, customers/create, customers/update, app/uninstalled")
bullet("Webhook HMAC verification enforced on every inbound request (SHA-256, Shopify-computed)")
bullet("Shopify API rate limits managed via token-bucket middleware in ingestion-svc (40 req/s Shopify Plus; 2 req/s basic)")
bullet("Storefront API used for product recommendations (read-only, public access token)")
bullet("ScriptTag injection for pixel (browse depth, add-to-cart signals) — single script tag, async, < 3 KB gzipped")

h2("6.5 WhatsApp Business API Integration")
bullet("Meta Cloud API (v18+) used directly for template management, message send, and delivery webhooks")
bullet("Template approval workflow: merchant submits template in dashboard → system auto-submits to Meta → approval status polled every 15 min")
bullet("Interactive message types supported: button (up to 3), list (up to 10 items), flow (multi-step)")
bullet("Opt-in collection: customer must explicitly opt in at checkout (Shopify checkout extension) before any WhatsApp message")
bullet("WABA (WhatsApp Business Account) provisioning: self-serve for brands with Meta Business Manager; white-glove for Enterprise tier")

h2("6.6 Security Architecture")
table_2col([
    ["Control", "Implementation"],
    ["Authentication",       "JWT (15-min access token) + refresh token (30-day, rotatable); Shopify session tokens for embedded app"],
    ["Authorisation",        "RBAC — roles: Owner, Admin, Analyst, Read-only; enforced at API gateway + service level"],
    ["Data isolation",       "Row-level security in PostgreSQL by shop_id; no cross-tenant data leakage possible"],
    ["Encryption at rest",   "AES-256 on RDS + S3; envelope encryption for PII fields (customer phone, email)"],
    ["Encryption in transit","TLS 1.3 minimum; HSTS enforced; mTLS between internal services"],
    ["PII handling",         "Phone / email stored as encrypted blobs; hashed in logs; 90-day retention then purge"],
    ["Rate limiting",        "Kong Gateway: 100 req/min per merchant API key; 10 req/min on auth endpoints"],
    ["Audit logging",        "All admin actions logged to immutable CloudWatch log group (90-day retention)"],
    ["Vulnerability mgmt",   "Snyk in CI/CD; weekly DAST scan (OWASP ZAP); annual pen-test"],
    ["Compliance",           "SOC 2 Type II (in progress); GDPR Article 28 DPA available; DPDP (India) compliant"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. DATA MODELS
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Core Data Models")

h2("7.1 Customer Record")
body("Primary entity representing a Shopify customer enriched with RetainIQ data.")
table_2col([
    ["Field", "Type / Notes"],
    ["id",                    "UUID (internal)"],
    ["shop_id",               "FK → shops.id (tenant isolation)"],
    ["shopify_customer_id",   "BigInt — Shopify native ID"],
    ["email_hash",            "SHA-256 of email — used for analytics; raw email encrypted separately"],
    ["phone_e164",            "Encrypted string — E.164 format"],
    ["churn_score",           "Float 0–100 — latest model output"],
    ["churn_score_updated_at","Timestamp"],
    ["predicted_reorder_date","Date per SKU — stored as JSONB {sku_id: date}"],
    ["ltv_to_date",           "Decimal — total spend"],
    ["order_count",           "Integer — lifetime orders"],
    ["last_order_at",         "Timestamp"],
    ["segment",               "Enum: champion / loyal / at_risk / dormant / new"],
    ["opt_in_whatsapp",       "Boolean — explicit consent"],
    ["opt_in_sms",            "Boolean"],
    ["opt_in_email",          "Boolean"],
    ["created_at",            "Timestamp"],
])

h2("7.2 Flow Execution Record")
table_2col([
    ["Field", "Type / Notes"],
    ["id",               "UUID"],
    ["customer_id",      "FK → customers.id"],
    ["flow_id",          "FK → flows.id"],
    ["trigger_event",    "Enum: order_paid / score_computed / delivery_confirmed / reorder_due"],
    ["step_id",          "FK → flow_steps.id"],
    ["channel",          "Enum: whatsapp / sms / email"],
    ["status",           "Enum: pending / sent / delivered / opened / clicked / converted / failed"],
    ["message_id",       "Provider message ID for delivery tracking"],
    ["scheduled_at",     "Timestamp — when message was scheduled"],
    ["sent_at",          "Timestamp"],
    ["delivered_at",     "Timestamp"],
    ["opened_at",        "Timestamp"],
    ["clicked_at",       "Timestamp"],
    ["converted_at",     "Timestamp — if purchase followed within attribution window"],
    ["revenue_attributed","Decimal — order value if converted"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. API SPECIFICATION (KEY ENDPOINTS)
# ══════════════════════════════════════════════════════════════════════════════
h1("8. API Specification (Key Endpoints)")

body("Base URL: https://api.retainiq.app/v1  |  Auth: Bearer JWT in Authorization header  |  Rate limit: 100 req/min")

h2("8.1 Customer Endpoints")
table_generic(
    ["Method", "Endpoint", "Description", "Response"],
    [
        ["GET",   "/customers/{id}",              "Fetch enriched customer record",           "200 Customer object"],
        ["GET",   "/customers/{id}/churn-score",  "Real-time churn score + feature breakdown","200 {score, features[]}"],
        ["GET",   "/customers/{id}/journey",      "Full post-purchase journey execution log", "200 FlowExecution[]"],
        ["POST",  "/customers/{id}/segments",     "Override AI segment",                      "200 updated Customer"],
        ["DELETE","/customers/{id}",              "GDPR erasure (purges all PII)",             "204 No Content"],
    ],
    col_widths=[Cm(1.5), Cm(6), Cm(5), Cm(2.5)]
)

h2("8.2 Flow Endpoints")
table_generic(
    ["Method", "Endpoint", "Description", "Response"],
    [
        ["GET",  "/flows",          "List all flows for merchant",           "200 Flow[]"],
        ["POST", "/flows",          "Create new flow",                       "201 Flow"],
        ["PUT",  "/flows/{id}",     "Update flow config",                    "200 Flow"],
        ["POST", "/flows/{id}/test","Send test message to merchant's number","200 {message_id}"],
        ["GET",  "/flows/{id}/analytics","Flow performance metrics",         "200 FlowAnalytics"],
    ],
    col_widths=[Cm(1.5), Cm(5), Cm(5.5), Cm(3)]
)

h2("8.3 Replenishment Endpoints")
table_generic(
    ["Method", "Endpoint", "Description", "Response"],
    [
        ["GET",  "/replenishment/forecast",         "30/60/90 day reorder forecast",                "200 ForecastDay[]"],
        ["GET",  "/customers/{id}/reorder-dates",   "Predicted reorder date per SKU for customer",  "200 {sku_id: date}[]"],
        ["POST", "/replenishment/manual-override",  "Override predicted date for customer×SKU",     "200 updated prediction"],
    ],
    col_widths=[Cm(1.5), Cm(6), Cm(5), Cm(2.5)]
)

h2("8.4 Webhooks (Outbound)")
body("RetainIQ fires outbound webhooks to merchant-configured URLs for key events.")
table_2col([
    ["Event", "Payload summary"],
    ["customer.score_updated",    "shop_id, customer_id, new_score, previous_score, top_features[]"],
    ["review.received",           "shop_id, customer_id, rating, text, has_ugc, sentiment"],
    ["reorder.placed",            "shop_id, customer_id, sku_id, order_id, attributed_flow_id"],
    ["flow.converted",            "shop_id, customer_id, flow_id, step_id, channel, revenue"],
    ["customer.churned",          "shop_id, customer_id, days_dormant, last_order_at"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. ML MODEL SPECIFICATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. ML Model Specifications")

h2("9.1 Churn Risk Model")
table_2col([
    ["Attribute", "Specification"],
    ["Model type",        "LightGBM Gradient-Boosted Decision Tree (classification)"],
    ["Target variable",   "Binary: did customer place a 2nd order within 90 days? (1 = yes, 0 = churn)"],
    ["Training data",     "36 months of anonymised Shopify order + browse data; min 10K orders per brand for brand-specific model"],
    ["Global model",      "Trained on pooled anonymised data from all brands (federated approach — no raw data sharing)"],
    ["Brand fine-tuning", "Transfer learning layer: brand-specific 500-order minimum to fine-tune top layer"],
    ["Feature importance","SHAP values computed per prediction; top-3 features returned in score API"],
    ["Validation",        "Stratified k-fold (k=5); AUC-ROC primary metric; Brier score secondary"],
    ["Target AUC",        "≥ 0.78 on held-out validation set"],
    ["Retraining",        "Global: weekly; Brand-specific: daily incremental via online learning"],
    ["Serving infra",     "Triton Inference Server on GPU node; model stored in MLflow registry"],
    ["Fallback",          "Rule-based score (AOV + category baseline) if ML service unavailable"],
])

h2("9.2 Replenishment Prediction Model")
table_2col([
    ["Attribute", "Specification"],
    ["Model type",        "Weibull-Gamma mixture model (Bayesian survival analysis)"],
    ["Target variable",   "Time-to-reorder in days for a customer×SKU pair"],
    ["Priors",            "Category-level priors from brand's historical order cadence"],
    ["Posterior update",  "Online Bayesian update on every new purchase event"],
    ["Uncertainty",       "95 % credible interval returned; used to set message timing"],
    ["Cold start",        "≥ 2 orders of same SKU required; category median until then"],
    ["MAE target",        "≤ 5 days for consumables (supplements, coffee, pet food, skincare)"],
    ["Explainability",    "Predicted date + \"why\" copy (\"Based on your 28-day usage pattern\") for WhatsApp message"],
])

h2("9.3 Review Sentiment Model")
table_2col([
    ["Attribute", "Specification"],
    ["Model type",        "Fine-tuned BERT-base (multilingual) for sentiment + topic classification"],
    ["Languages",         "English, Hindi, Tamil, Indonesian, Brazilian Portuguese"],
    ["Classes",           "Sentiment: positive / neutral / negative; Topic: quality / delivery / size / value"],
    ["Use case",          "Classify 3-star reviews to route to correct support queue; flag product issues"],
    ["Latency",           "< 200 ms per review (async, non-blocking)"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. PERFORMANCE & SLA REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Performance & SLA Requirements")

h2("10.1 Latency Targets")
table_generic(
    ["Operation", "p50", "p95", "p99"],
    [
        ["Shopify webhook receive → Kafka produce",    "< 50 ms",  "< 150 ms",  "< 300 ms"],
        ["Churn score computation",                    "< 80 ms",  "< 200 ms",  "< 500 ms"],
        ["WhatsApp message dispatch",                  "< 500 ms", "< 2 s",     "< 5 s"],
        ["Branded tracking page (LCP)",                "< 1.2 s",  "< 2.0 s",   "< 3.5 s"],
        ["Dashboard data load (7-day view)",           "< 800 ms", "< 2 s",     "< 4 s"],
        ["Churn score API (external)",                 "< 80 ms",  "< 200 ms",  "< 300 ms"],
    ],
    col_widths=[Cm(6), Cm(2.5), Cm(2.5), Cm(2.5)]
)

h2("10.2 Availability SLA")
table_2col([
    ["Tier", "Availability SLA"],
    ["Starter",    "99.5 % (planned maintenance excluded)"],
    ["Growth",     "99.9 % (< 8.7 hrs/year downtime)"],
    ["Enterprise", "99.95 % with dedicated support SLA"],
])

h2("10.3 Scale Targets")
table_2col([
    ["Metric", "Target"],
    ["Peak webhook ingestion",    "50 000 events/minute (Black Friday / Diwali sale spikes)"],
    ["Active customers tracked",  "50 M across all tenants"],
    ["Messages dispatched",       "5 M / day across channels"],
    ["Concurrent dashboard users","10 000 concurrent merchant sessions"],
    ["Data retention",            "24 months hot; 5 years cold (S3 Glacier)"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. MONETISATION MODEL
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Monetisation Model")

h2("11.1 Pricing Tiers")
table_generic(
    ["", "Starter", "Growth", "Enterprise"],
    [
        ["Price",                   "$99/mo",           "$299/mo",                 "Custom (from $999/mo)"],
        ["Active customers",        "Up to 5 000",      "Up to 50 000",            "Unlimited"],
        ["WhatsApp messages",       "1 000/mo incl.",   "10 000/mo incl.",         "Custom pool"],
        ["Email messages",          "Unlimited",        "Unlimited",               "Unlimited"],
        ["SMS messages",            "Pay-as-you-go",    "2 000/mo incl.",          "Custom pool"],
        ["Churn scoring",           "Yes",              "Yes",                     "Yes + custom model"],
        ["Replenishment AI",        "No",               "Yes",                     "Yes + custom thresholds"],
        ["A/B testing",             "No",               "Yes",                     "Yes + multivariate"],
        ["BigQuery / Snowflake",    "No",               "No",                      "Yes"],
        ["API access",              "Read-only",        "Full",                    "Full + higher rate limits"],
        ["Branded tracking domain", "retainiq subdomain","Custom domain",           "Custom domain + SSL mgmt"],
        ["Support SLA",             "Email (48 h)",     "Email + chat (24 h)",     "Dedicated CSM + < 4 h"],
        ["Uptime SLA",              "99.5 %",           "99.9 %",                  "99.95 %"],
    ],
    col_widths=[Cm(4), Cm(3.5), Cm(3.5), Cm(4)]
)

h2("11.2 Usage-Based Overages")
table_2col([
    ["Resource", "Overage Price"],
    ["WhatsApp messages (over plan)",  "$0.04 / message (carrier surcharge pass-through + margin)"],
    ["SMS messages (over plan)",       "$0.025 / message (US/UK); $0.01 / message (India)"],
    ["Active customers (over plan)",   "$0.005 / customer / month"],
    ["Additional data exports",        "$0.10 / 1 000 rows"],
])

h2("11.3 Revenue Projections")
table_generic(
    ["Month", "Installs", "MRR", "ARR Run Rate"],
    [
        ["Month 3",   "200",    "$28 500",   "$342 000"],
        ["Month 6",   "650",    "$92 000",   "$1.1 M"],
        ["Month 9",   "1 300",  "$182 000",  "$2.18 M"],
        ["Month 12",  "2 000",  "$200 000",  "$2.4 M"],
    ],
    col_widths=[Cm(2.5), Cm(3), Cm(4), Cm(5)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. COMPETITIVE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Competitive Analysis")

h2("12.1 Competitive Landscape")
table_generic(
    ["Competitor",  "Strength",                        "Gap vs RetainIQ Pro",                       "Our counter"],
    [
        ["Klaviyo",     "Best-in-class email flows",       "No WhatsApp; no churn scoring; no replenishment AI","Omnichannel + AI layer on top or instead of"],
        ["Yotpo",       "Strong review + loyalty platform","No sequencing; no churn; no WhatsApp",             "Review loop is one feature, not whole product"],
        ["Postscript",  "SMS for Shopify (US-focused)",    "SMS only; no WhatsApp; no scoring",                "WhatsApp-first (87% open rate) + channel agnostic"],
        ["AfterShip",   "Best tracking page",              "No retention; no AI; no messaging",                "Branded tracking is 1 pillar of our 4-pillar platform"],
        ["Gorgias",     "Best post-purchase support",      "Reactive (support), not proactive (retention)",    "We pre-empt issues via churn scoring + review routing"],
        ["Retention.com","Email capture + retargeting",    "Acquisition focus; no post-purchase journey",      "We're post-purchase specialists; no overlap"],
        ["Bird (SparkPost)","Omnichannel messaging infra",  "Infrastructure only; no Shopify native; no AI",   "Turnkey Shopify app; AI included; no dev required"],
    ],
    col_widths=[Cm(3), Cm(3.5), Cm(4.5), Cm(4)]
)

h2("12.2 Differentiation Matrix")
table_generic(
    ["Capability",             "RetainIQ", "Klaviyo", "Yotpo", "Postscript", "AfterShip"],
    [
        ["Churn scoring (AI)",          "Yes",  "No",   "No",  "No",   "No"],
        ["WhatsApp Business (native)",  "Yes",  "No",   "No",  "No",   "No"],
        ["SMS + Email + WA unified",    "Yes",  "Email","No",  "SMS",  "No"],
        ["Replenishment AI",            "Yes",  "No",   "No",  "No",   "No"],
        ["UGC capture + routing",       "Yes",  "No",   "Yes", "No",   "No"],
        ["Branded tracking page",       "Yes",  "No",   "No",  "No",   "Yes"],
        ["One-tap WhatsApp reorder",    "Yes",  "No",   "No",  "No",   "No"],
        ["NPS engine",                  "Yes",  "No",   "Yes", "No",   "No"],
        ["Shopify App Store native",    "Yes",  "Yes",  "Yes", "Yes",  "Yes"],
        ["< 1 hr setup",                "Yes",  "No",   "No",  "Yes",  "Yes"],
    ],
    col_widths=[Cm(5.5), Cm(2), Cm(2), Cm(2), Cm(2.5), Cm(2)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. GO-TO-MARKET STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Go-to-Market Strategy")

h2("13.1 Launch Phases")

h3("Phase 1 — Beta (Months 0–3): 50 Design Partners")
bullet("Hand-pick 50 D2C brands in consumables / beauty categories from founder networks")
bullet("Free for 90 days in exchange for: weekly feedback calls, case study rights, referrals")
bullet("Success criterion: 80 % report measurable repeat-rate improvement; 3 publishable case studies")
bullet("Primary goal: product-market fit evidence and retention metric benchmarks")

h3("Phase 2 — Shopify App Store Launch (Months 3–6)")
bullet("List on Shopify App Store (retention / loyalty category)")
bullet("Shopify App Store optimisation: video demo, 50+ reviews seeded from beta cohort")
bullet("Content marketing: SEO content targeting 'Shopify retention app', 'reduce D2C churn', 'WhatsApp for Shopify'")
bullet("Founder community outreach: YC alumni groups, Founder's Network India, Slack communities (DTCX, Operators)")
bullet("Target: 200 installs, $28K MRR by end of Month 3 post-launch")

h3("Phase 3 — Growth (Months 6–12)")
bullet("Partner programme: Klaviyo agency partners, Shopify Plus partners, WhatsApp BSPs")
bullet("WhatsApp geo expansion: India, Indonesia, Brazil — localised templates, local pricing")
bullet("Direct sales: outbound to Shopify Plus merchants ($5M+ ARR) for Enterprise tier")
bullet("Referral programme: merchants earn 2 months free per referred install (from Growth tier)")
bullet("Integrations marketplace: Gorgias, LoyaltyLion, Recharge, Okendo — native bi-directional integrations")

h2("13.2 Acquisition Channels & CAC Targets")
table_generic(
    ["Channel", "Target CAC", "Volume/mo (Month 12)", "Notes"],
    [
        ["Shopify App Store (organic)", "$0",    "600 installs",  "SEO + rating flywheel"],
        ["Content / SEO",               "$120",  "300 installs",  "LTV:CAC > 12:1 at $299/mo"],
        ["Partner / agency referral",   "$80",   "400 installs",  "Revenue-share: 20 % first year"],
        ["Paid (Google/Meta)",          "$250",  "400 installs",  "Growth+ tier only (higher LTV)"],
        ["Direct sales (Enterprise)",   "$800",  "50 accounts",   "High LTV ($1K+/mo) justifies cost"],
    ],
    col_widths=[Cm(4.5), Cm(2.5), Cm(3.5), Cm(4.5)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. PRODUCT ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Product Roadmap")

h2("Phase 1 — Foundation (Months 1–3)")
table_generic(
    ["Epic", "Features", "Priority", "Owner"],
    [
        ["Shopify Integration",      "OAuth install, webhook registration, pixel injection, order sync",         "P0", "Backend"],
        ["Churn Scoring V1",         "Rule-based cold-start scoring, LightGBM global model, score API",          "P0", "ML"],
        ["Flow Engine V1",           "Journey builder UI, 5 pre-built templates, WhatsApp + Email dispatch",     "P0", "Product"],
        ["Branded Tracking Page V1", "Carrier integration (top 10), customisable layout, basic analytics",       "P0", "Frontend"],
        ["Review Capture V1",        "T+72h WhatsApp review request, 5-star routing, storefront widget",         "P0", "Product"],
        ["Merchant Dashboard V1",    "Retention overview, flow performance, channel analytics",                  "P0", "Frontend"],
        ["Billing & Subscriptions",  "Stripe integration, plan enforcement, usage metering",                     "P0", "Backend"],
    ],
    col_widths=[Cm(3.5), Cm(7), Cm(1.5), Cm(2)]
)

h2("Phase 2 — Intelligence (Months 4–6)")
table_generic(
    ["Epic", "Features", "Priority", "Owner"],
    [
        ["Replenishment AI V1",    "Bayesian model, reorder date API, WhatsApp one-tap reorder flow",       "P0", "ML + Product"],
        ["Churn Scoring V2",       "Brand fine-tuning, SHAP feature explanations, model drift monitoring",  "P0", "ML"],
        ["SMS Channel",            "Twilio integration, country routing, opt-in management",                "P0", "Backend"],
        ["A/B Testing",            "Split traffic, significance detection, auto winner promotion",          "P1", "Product"],
        ["NPS Engine",             "T+14d NPS survey, promoter→referral pipeline, detractor→support",      "P1", "Product"],
        ["UGC Library",            "Media upload, approval workflow, Meta republish integration",           "P1", "Frontend"],
        ["Customer Segments",      "AI segment auto-classification, Klaviyo sync, CSV export",              "P1", "ML + Backend"],
    ],
    col_widths=[Cm(3.5), Cm(7), Cm(1.5), Cm(2)]
)

h2("Phase 3 — Scale (Months 7–9)")
table_generic(
    ["Epic", "Features", "Priority", "Owner"],
    [
        ["WhatsApp Geo Expansion",  "India / Indonesia / Brazil localisation, local BSP fallback",         "P0", "Backend"],
        ["Enterprise Features",     "SSO (SAML/OIDC), BigQuery sync, dedicated onboarding, SLA",          "P0", "Platform"],
        ["Integrations",            "Gorgias, Recharge, LoyaltyLion, Okendo bi-directional integrations", "P1", "Partnerships"],
        ["Review AI",               "Multilingual sentiment + topic classification, issue alerting",       "P1", "ML"],
        ["Referral Engine",         "Per-customer referral links, reward tracking, fraud detection",       "P1", "Product"],
        ["Predictive LTV",          "12-month LTV forecast per customer with confidence interval",         "P2", "ML"],
    ],
    col_widths=[Cm(3.5), Cm(7), Cm(1.5), Cm(2)]
)

h2("Phase 4 — Moat (Months 10–12)")
table_generic(
    ["Epic", "Features", "Priority", "Owner"],
    [
        ["Loyalty Engine",          "Points, tiers, perks — integrated with replenishment + review rewards","P1", "Product"],
        ["Subscription Intelligence","Predict subscription churn 30 days out; rescue flows",               "P1", "ML + Product"],
        ["AI Flow Builder",         "Natural language flow creation: \"build a flow for high-risk customers in beauty\"","P2","AI"],
        ["WhatsApp Commerce",       "Product catalogue in WA, in-thread checkout, payment collection",     "P1", "Platform"],
        ["Global Expansion",        "Stripe multi-currency billing; EU data residency (Frankfurt region)", "P0", "Platform"],
        ["SOC 2 Type II Audit",     "Evidence collection, auditor engagement, report publication",         "P0", "Security"],
    ],
    col_widths=[Cm(3.5), Cm(7), Cm(1.5), Cm(2)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. SUCCESS METRICS & KPIs
# ══════════════════════════════════════════════════════════════════════════════
h1("15. Success Metrics & KPIs")

h2("15.1 North Star Metric")
callout("North Star", "RetainIQ-attributed repeat purchases per month (across all merchant installs). "
        "This is the single metric that aligns product, engineering, and go-to-market.")

h2("15.2 L1 Product KPIs (Reviewed Monthly)")
table_generic(
    ["KPI", "Definition", "Month 6 Target", "Month 12 Target"],
    [
        ["Active installs",          "Shopify stores with ≥ 1 flow active in last 30 days", "650",         "2 000"],
        ["MRR",                      "Monthly recurring revenue",                            "$92K",        "$200K"],
        ["Net Revenue Retention",    "(MRR end - churn + expansion) / MRR start",            "> 95 %",      "> 105 %"],
        ["Merchant repeat rate lift","Avg (post-RetainIQ repeat rate - pre-RetainIQ)",        "> 25 %",      "> 34 %"],
        ["Flow active rate",         "% installs with ≥ 1 live flow",                        "> 80 %",      "> 85 %"],
        ["WA message open rate",     "Platform-wide WhatsApp open rate",                     "> 80 %",      "> 85 %"],
        ["Review collection rate",   "Reviews received / review requests sent",              "> 18 %",      "> 22 %"],
        ["Reorder conversion rate",  "Reorders placed / replenishment nudges sent",          "> 15 %",      "> 18 %"],
    ],
    col_widths=[Cm(4.5), Cm(5), Cm(2.5), Cm(3)]
)

h2("15.3 L2 Engineering KPIs (Reviewed Weekly)")
table_2col([
    ["KPI", "Target"],
    ["Webhook processing latency p99",   "< 300 ms"],
    ["Score computation p99",            "< 500 ms"],
    ["Message delivery success rate",    "≥ 99.5 %"],
    ["Platform uptime (Growth tier)",    "≥ 99.9 %"],
    ["CI/CD deploy frequency",           "≥ 3 deployments / week"],
    ["Mean time to recovery (MTTR)",     "< 30 min for P0 incidents"],
    ["Error rate (5xx on public APIs)",  "< 0.1 %"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 16. RISK ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1("16. Risk Analysis & Mitigations")

table_generic(
    ["Risk", "Probability", "Impact", "Mitigation"],
    [
        ["Meta WhatsApp API policy change\n(template rejection, pricing hike)",
         "Medium", "High",
         "Multi-BSP fallback (Gupshup, Twilio, Meta direct). SMS fallback always active. T&C monitoring."],
        ["Shopify Partner Agreement violation\n(data handling, revenue share change)",
         "Low", "Critical",
         "Legal review of Partner Agreement quarterly. Data handling in compliance with Shopify data practices."],
        ["ML model underperformance\n(low-data brands, new categories)",
         "Medium", "Medium",
         "Rule-based fallback always active. Brand-level performance monitoring with auto-fallback trigger."],
        ["Competitor (Klaviyo) adds WhatsApp\nand churn scoring",
         "Medium", "High",
         "Speed of execution + deeper Shopify native integration + replenishment AI as moat. Patent filing for core algorithms."],
        ["Enterprise data breach / PII leak",
         "Low", "Critical",
         "SOC 2, pen-testing, encryption at field level, row-level security, no raw PII in logs, cyber insurance."],
        ["Merchant churn due to poor\nonboarding / no quick wins",
         "High", "High",
         "< 1 hr setup mandate. Pre-built flows with guaranteed 30-day result or refund. Dedicated onboarding for Growth+."],
        ["WhatsApp message spam complaints\nlead to account ban",
         "Medium", "Critical",
         "Strict opt-in only. Quality rating monitoring via Meta API. Auto-pause if quality drops below 'Medium'."],
        ["Multi-currency / tax complexity\nfor global expansion",
         "Low", "Medium",
         "Stripe multi-currency + Avalara tax automation. Regional pricing in Phase 4."],
    ],
    col_widths=[Cm(4), Cm(2), Cm(1.5), Cm(7.5)]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 17. COMPLIANCE & DATA GOVERNANCE
# ══════════════════════════════════════════════════════════════════════════════
h1("17. Compliance & Data Governance")

h2("17.1 Regulatory Compliance")
table_2col([
    ["Regulation", "Implementation"],
    ["GDPR (EU)",                "DPA available; right to erasure implemented (DELETE /customers/{id}); data residency option (EU-West)"],
    ["CCPA (California)",        "Opt-out of data sale link on tracking pages; PII deletion on request within 45 days"],
    ["DPDP 2023 (India)",        "Consent-first WhatsApp opt-in; data principal rights implemented; data localisation option"],
    ["TCPA (US SMS)",            "Double opt-in for SMS; time-of-day restrictions (8am–9pm local); carrier compliance"],
    ["WhatsApp Commerce Policy", "Prohibited category checks on install; template pre-approval; quality rating monitoring"],
    ["SOC 2 Type II",            "In-progress; Type I expected Month 6; Type II expected Month 12"],
    ["PCI DSS",                  "No card data stored; Stripe handles all card processing; SAQ-A scope only"],
])

h2("17.2 Data Retention Policy")
table_2col([
    ["Data type", "Retention"],
    ["Customer PII (email, phone)",    "Retained while merchant account active; purged 30 days after account closure"],
    ["Order data",                     "24 months hot (PostgreSQL); 5 years cold (S3 Glacier)"],
    ["Message content",                "Message templates stored indefinitely; individual sends stored 12 months"],
    ["ML training data",               "Anonymised aggregates only; no PII in training sets"],
    ["Audit logs",                     "90 days in CloudWatch; 1 year in S3"],
    ["UGC assets",                     "Retained while merchant account active + 30-day grace"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 18. APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
h1("18. Appendix")

h2("A. Glossary")
table_2col([
    ["Term", "Definition"],
    ["AOV",           "Average Order Value"],
    ["BFCM",          "Black Friday / Cyber Monday — peak Shopify sales period"],
    ["BSP",           "Business Solution Provider — WhatsApp API reseller"],
    ["CAC",           "Customer Acquisition Cost"],
    ["CTA",           "Call to Action"],
    ["D2C",           "Direct to Consumer"],
    ["LTV",           "Lifetime Value — total revenue from a customer over their relationship"],
    ["MAE",           "Mean Absolute Error — ML model accuracy metric"],
    ["NPS",           "Net Promoter Score — customer loyalty metric"],
    ["NRR",           "Net Revenue Retention — expansion minus churn over period"],
    ["PII",           "Personally Identifiable Information"],
    ["ROAS",          "Return on Ad Spend"],
    ["SKU",           "Stock Keeping Unit — individual product variant"],
    ["UGC",           "User-Generated Content (photos, videos, reviews)"],
    ["WABA",          "WhatsApp Business Account — Meta-issued business entity for WhatsApp API"],
])

h2("B. Key Third-Party Dependencies")
table_generic(
    ["Dependency", "Purpose", "Tier", "Fallback"],
    [
        ["Meta WhatsApp Cloud API",  "WhatsApp message delivery",        "All",          "Gupshup BSP"],
        ["Twilio",                   "SMS delivery (global)",             "All",          "Vonage"],
        ["MSG91",                    "SMS delivery (India)",              "All",          "Twilio"],
        ["SendGrid",                 "Transactional email",               "All",          "SES"],
        ["EasyPost",                 "Carrier tracking aggregation",      "All",          "AfterShip API"],
        ["Stripe",                   "Billing and subscriptions",         "All",          "None (critical)"],
        ["AWS (EKS, RDS, MSK, S3)", "Core infrastructure",               "All",          "Multi-AZ HA only"],
        ["Shopify",                  "Merchant store data, webhooks",     "All",          "None (core platform)"],
        ["Google Business Profile",  "Auto-publish 5★ reviews",          "All",          "Manual publish"],
        ["Meta Graph API",           "Facebook / Instagram review publish","All",         "Manual publish"],
        ["MLflow",                   "ML experiment tracking, model store","Internal",    "S3 + version tags"],
    ],
    col_widths=[Cm(3.5), Cm(4), Cm(2), Cm(4.5)]
)

h2("C. Open Questions for Engineering Review")
bullet("Q1: Should replenishment model run per-brand or as a global multi-task model? (Brand isolation vs data sparsity tradeoff)")
bullet("Q2: WhatsApp conversation-based pricing ($0.0147/conv in India) — how do we meter and pass through without surprising merchants?")
bullet("Q3: Shopify Flow (native automation) — build native extension or keep as standalone app? Impacts distribution significantly.")
bullet("Q4: WABA provisioning — self-serve (longer setup) vs. RetainIQ-managed shared WABA (faster, but limits brand identity on sender profile)?")
bullet("Q5: ML training data pooling — explicit merchant consent required? Legal review needed before cross-tenant model training.")

# ── Final page ────────────────────────────────────────────────────────────────
doc.add_page_break()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_p.paragraph_format.space_before = Pt(80)
add_run(end_p, "— End of Document —", italic=True, size=12, color=MID_TEXT)

doc.add_paragraph()
ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(ver_p,
        f"RetainIQ Pro PRD v1.0  |  {datetime.date.today().strftime('%B %d, %Y')}  |  Confidential",
        size=10, color=MID_TEXT)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/sayujpillai/Desktop/retainIQ/RetainIQ-Pro-PRD.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
